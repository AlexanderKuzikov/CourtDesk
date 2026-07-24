#!/usr/bin/env python3
"""Convert CourtDesk markdown docs to DOCX and HTML (for PDF pipeline)."""

import re
import sys
import os
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

BASE = Path(__file__).parent

def read_md(filename: str) -> str:
    path = BASE / filename
    with open(path, 'r', encoding='utf-8') as f:
        return f.read()

def md_to_docx(md_text: str, title: str, output_name: str):
    """Parse a subset of markdown and build a DOCX."""
    doc = Document()

    # --- Page setup ---
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.0)

    # --- Styles ---
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10.5)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    # --- Title page ---
    for _ in range(6):
        doc.add_paragraph('')

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(title)
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x1e, 0x29, 0x3b)

    doc.add_paragraph('')
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run('CourtDesk — CRM-оркестратор поиска и мониторинга судебных дел РФ')
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x64, 0x74, 0x8b)

    doc.add_paragraph('')
    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_p.add_run('2026-07-24')
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x94, 0xa3, 0xb8)

    doc.add_page_break()

    # --- Parse body ---
    lines = md_text.split('\n')
    i = 0
    in_code_block = False
    code_buffer = []
    in_table = False
    table_buffer = []
    in_list = False
    list_items = []

    def flush_code():
        nonlocal code_buffer
        if not code_buffer:
            return
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run('\n'.join(code_buffer))
        run.font.name = 'Consolas'
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        # Add shading
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F1F5F9"/>')
        p.paragraph_format.element.get_or_add_pPr().append(shading)
        code_buffer = []

    def flush_table():
        nonlocal table_buffer
        if not table_buffer:
            return
        rows_data = []
        for line in table_buffer:
            cells = [c.strip() for c in line.split('|')[1:-1]]
            rows_data.append(cells)

        if len(rows_data) < 2:
            table_buffer = []
            return

        # Determine columns
        ncols = max(len(r) for r in rows_data)
        if ncols < 2:
            table_buffer = []
            return

        table = doc.add_table(rows=len(rows_data), cols=ncols)
        table.style = 'Light Grid Accent 1'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for ri, row_data in enumerate(rows_data):
            for ci in range(ncols):
                val = row_data[ci] if ci < len(row_data) else ''
                cell = table.cell(ri, ci)
                cell.text = val
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_before = Pt(1)
                    paragraph.paragraph_format.space_after = Pt(1)
                    for run in paragraph.runs:
                        run.font.size = Pt(9)
                        if ri == 0:
                            run.bold = True

        doc.add_paragraph('')  # spacing
        table_buffer = []

    def flush_list():
        nonlocal list_items
        if not list_items:
            return
        for item in list_items:
            p = doc.add_paragraph(style='List Bullet')
            p.text = item
            p.paragraph_format.space_after = Pt(2)
        list_items = []

    while i < len(lines):
        line = lines[i]

        # Code block
        if line.startswith('```'):
            if in_code_block:
                in_code_block = False
                flush_code()
            else:
                in_code_block = True
                code_buffer = []
            i += 1
            continue

        if in_code_block:
            code_buffer.append(line)
            i += 1
            continue

        # Skip horizontal rules
        if line.strip() in ('---', '***', '___'):
            i += 1
            continue

        # Table detection
        if line.strip().startswith('|') and line.strip().endswith('|'):
            # Check if it's a separator row
            if re.match(r'^\|[\s\-:]+\|$', line.strip()):
                i += 1
                continue
            flush_list()
            table_buffer.append(line.strip())
            i += 1
            continue
        else:
            if table_buffer:
                flush_table()

        # Headings
        heading_match = re.match(r'^(#{1,4})\s+(.+)$', line)
        if heading_match:
            flush_list()
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()
            h = doc.add_heading(text, level=min(level, 4))
            h.paragraph_format.space_before = Pt(14 if level <= 2 else 8)
            h.paragraph_format.space_after = Pt(4)
            i += 1
            continue

        # Unordered list
        list_match = re.match(r'^[\s]*[-*+]\s+(.+)$', line)
        if list_match:
            list_items.append(list_match.group(1))
            i += 1
            continue
        else:
            flush_list()

        # Ordered list
        olist_match = re.match(r'^[\s]*\d+[.)]\s+(.+)$', line)
        if olist_match:
            list_items.append(olist_match.group(1))
            i += 1
            continue
        else:
            flush_list()

        # Empty line
        if not line.strip():
            i += 1
            continue

        # Regular paragraph with inline formatting
        text = line.strip()

        # Skip badges like ✅  but keep their text
        text = re.sub(r'\*\*', '', text)  # remove ** markers

        # Handle bold markers **text**
        parts = re.split(r'(\*\*.*?\*\*)', text)
        p = doc.add_paragraph()
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                run.bold = True
            else:
                # Process inline code
                subparts = re.split(r'(`[^`]+`)', part)
                for sp in subparts:
                    if sp.startswith('`') and sp.endswith('`'):
                        r = p.add_run(sp[1:-1])
                        r.font.name = 'Consolas'
                        r.font.size = Pt(9)
                        r.font.color.rgb = RGBColor(0xE5, 0x3E, 0x3E)
                    else:
                        p.add_run(sp)

        p.paragraph_format.space_after = Pt(4)
        i += 1

    # Flush remaining
    if code_buffer:
        flush_code()
    if table_buffer:
        flush_table()
    if list_items:
        flush_list()

    # --- Save ---
    output_path = BASE / output_name
    doc.save(str(output_path))
    print(f'[DOCX] Created: {output_path.name} ({output_path.stat().st_size / 1024:.0f} KB)')
    return output_path


def md_to_html(md_text: str, title: str, output_name: str):
    """Convert markdown to standalone HTML for Puppeteer PDF pipeline."""
    lines = md_text.split('\n')
    html_parts = []
    i = 0
    in_code = False
    in_table = False
    table_rows = []
    in_list = False
    in_ol = False
    list_tag = 'ul'

    def esc(s):
        return (s.replace('&', '&amp;')
                 .replace('<', '&lt;')
                 .replace('>', '&gt;'))

    def flush_table():
        nonlocal table_rows, in_table
        if not table_rows:
            return
        html_parts.append('<table>')
        for ri, row in enumerate(table_rows):
            tag = 'th' if ri == 0 else 'td'
            html_parts.append(f'<tr>{"".join(f"<{tag}>{esc(c.strip())}</{tag}>" for c in row.split("|")[1:-1])}</tr>')
        html_parts.append('</table>')
        table_rows = []
        in_table = False

    def flush_list():
        nonlocal in_list, in_ol
        if in_list or in_ol:
            html_parts.append(f'</{list_tag}>')
            in_list = False
            in_ol = False

    while i < len(lines):
        line = lines[i]

        if line.startswith('```'):
            if in_code:
                html_parts.append('</code></pre>')
                in_code = False
            else:
                html_parts.append('<pre><code>')
                in_code = True
            i += 1
            continue

        if in_code:
            html_parts.append(esc(line) + '\n')
            i += 1
            continue

        if line.strip() in ('---', '***', '___'):
            flush_list()
            flush_table()
            html_parts.append('<hr>')
            i += 1
            continue

        # Table
        if line.strip().startswith('|') and line.strip().endswith('|'):
            if re.match(r'^\|[\s\-:]+\|$', line.strip()):
                i += 1
                continue
            flush_list()
            in_table = True
            table_rows.append(line.strip())
            i += 1
            continue
        else:
            flush_table()

        # Headings
        hm = re.match(r'^(#{1,4})\s+(.+)$', line)
        if hm:
            flush_list()
            lvl = len(hm.group(1))
            html_parts.append(f'<h{lvl}>{esc(hm.group(2).strip())}</h{lvl}>')
            i += 1
            continue

        # Lists
        lm = re.match(r'^[\s]*[-*+]\s+(.+)$', line)
        if lm:
            if not in_list:
                flush_list()
                in_list = True
                html_parts.append('<ul>')
            html_parts.append(f'<li>{esc(lm.group(1))}</li>')
            i += 1
            continue
        olm = re.match(r'^[\s]*\d+[.)]\s+(.+)$', line)
        if olm:
            if not in_ol:
                flush_list()
                in_ol = True
                html_parts.append('<ol>')
            html_parts.append(f'<li>{esc(olm.group(1))}</li>')
            i += 1
            continue
        flush_list()

        # Empty
        if not line.strip():
            i += 1
            continue

        # Paragraph — handle inline formatting
        text = esc(line.strip())
        text = re.sub(r'\*\*(.+?)\*\*', r'<strong>\1</strong>', text)
        text = re.sub(r'`([^`]+)`', r'<code>\1</code>', text)
        # Handle [text](url)
        text = re.sub(r'\[([^\]]+)\]\(([^)]+)\)', r'<a href="\2">\1</a>', text)
        html_parts.append(f'<p>{text}</p>')
        i += 1

    flush_table()
    flush_list()
    if in_code:
        html_parts.append('</code></pre>')

    body = '\n'.join(html_parts)

    full_html = f'''<!DOCTYPE html>
<html lang="ru">
<head>
<meta charset="UTF-8">
<title>{esc(title)}</title>
<style>
  @page {{ size: A4; margin: 2cm 2.5cm 2cm 2cm; }}
  body {{ font-family: 'Segoe UI', Calibri, sans-serif; font-size: 11pt; line-height: 1.5; color: #1e293b; }}
  h1 {{ font-size: 18pt; color: #1e293b; border-bottom: 2px solid #38bdf8; padding-bottom: 4px; margin-top: 28px; }}
  h2 {{ font-size: 14pt; color: #1e293b; margin-top: 22px; }}
  h3 {{ font-size: 12pt; color: #475569; margin-top: 16px; }}
  h4 {{ font-size: 11pt; color: #64748b; margin-top: 12px; }}
  p {{ margin: 4px 0 8px; }}
  table {{ width: 100%; border-collapse: collapse; margin: 8px 0; font-size: 9.5pt; }}
  th {{ background: #1e293b; color: white; padding: 5px 8px; text-align: left; font-weight: 600; }}
  td {{ padding: 4px 8px; border: 1px solid #cbd5e1; }}
  tr:nth-child(even) td {{ background: #f8fafc; }}
  code {{ font-family: 'Consolas', 'Cascadia Code', monospace; font-size: 9pt; background: #f1f5f9; padding: 1px 4px; border-radius: 3px; color: #e53e3e; }}
  pre {{ background: #f1f5f9; padding: 10px 14px; border-radius: 6px; border: 1px solid #e2e8f0; font-size: 9pt; overflow-x: auto; }}
  pre code {{ background: none; padding: 0; color: #1e293b; }}
  ul, ol {{ margin: 4px 0 8px; padding-left: 24px; }}
  li {{ margin: 2px 0; }}
  hr {{ border: none; border-top: 2px solid #e2e8f0; margin: 16px 0; }}
  a {{ color: #2563eb; text-decoration: none; }}
  strong {{ color: #0f172a; }}
  .cover {{ text-align: center; padding-top: 180px; }}
  .cover h1 {{ font-size: 24pt; border: none; color: #0f172a; margin-bottom: 8px; }}
  .cover .sub {{ font-size: 13pt; color: #64748b; }}
  .cover .date {{ font-size: 11pt; color: #94a3b8; margin-top: 16px; }}
</style>
</head>
<body>
<div class="cover">
  <h1>{esc(title)}</h1>
  <p class="sub">CourtDesk — CRM-оркестратор поиска и мониторинга судебных дел РФ</p>
  <p class="date">2026-07-24</p>
</div>
<div style="page-break-before: always;"></div>
{body}
</body>
</html>'''

    output_path = BASE / output_name
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(full_html)
    print(f'[HTML] Created: {output_path.name} ({output_path.stat().st_size / 1024:.0f} KB)')
    return output_path


def main():
    files = [
        ('CRM-INTEGRATION.md', 'Интеграция CourtDesk с 1С CRM'),
        ('WEBUI-FOR-LAWYER.md', 'CourtDesk Web UI — для юриста'),
    ]

    for md_file, title in files:
        text = read_md(md_file)
        # DOCX
        docx_name = md_file.replace('.md', '.docx')
        md_to_docx(text, title, docx_name)
        # HTML (for PDF pipeline)
        html_name = md_file.replace('.md', '.html')
        md_to_html(text, title, html_name)


if __name__ == '__main__':
    main()
